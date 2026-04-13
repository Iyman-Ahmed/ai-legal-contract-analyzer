"""
Document parser for legal contracts.

Supports PDF, DOCX, and plain text.
Primary: PyMuPDF (fitz) for PDFs — fast, structure-preserving.
Fallback: pdfplumber for complex multi-column layouts.
DOCX: python-docx with heading hierarchy preserved.
"""

import re
import io
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, Union
import logging

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    """Represents a single page from a document."""
    page_number: int
    text: str
    headings: list[str] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Full parsed document with structure metadata."""
    filename: str
    file_type: str
    pages: list[ParsedPage]
    full_text: str
    total_pages: int
    metadata: dict = field(default_factory=dict)

    def get_text_by_page(self, page_num: int) -> str:
        for p in self.pages:
            if p.page_number == page_num:
                return p.text
        return ""


class DocumentParser:
    """
    Parses legal documents (PDF, DOCX, TXT) into structured text.
    Uses PyMuPDF as primary parser, pdfplumber as fallback.
    """

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}

    def parse(self, file_path: Union[str, Path], file_bytes: Optional[bytes] = None) -> ParsedDocument:
        """
        Parse a document from path or raw bytes.

        Args:
            file_path: Path to the document (used for format detection)
            file_bytes: Raw file bytes (used when file comes from Gradio upload)

        Returns:
            ParsedDocument with text and structure metadata
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format '{ext}'. Supported: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        if file_bytes is None:
            with open(path, "rb") as f:
                file_bytes = f.read()

        if len(file_bytes) > 20 * 1024 * 1024:  # 20 MB
            raise ValueError("File exceeds 20 MB limit. Please upload a smaller document.")

        logger.info(f"Parsing {path.name} ({ext}, {len(file_bytes)/1024:.1f} KB)")

        if ext == ".pdf":
            return self._parse_pdf(file_bytes, path.name)
        elif ext == ".docx":
            return self._parse_docx(file_bytes, path.name)
        else:
            return self._parse_txt(file_bytes, path.name)

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse PDF using PyMuPDF; fall back to pdfplumber on failure."""
        try:
            return self._parse_pdf_pymupdf(file_bytes, filename)
        except Exception as e:
            logger.warning(f"PyMuPDF failed ({e}), falling back to pdfplumber")
            return self._parse_pdf_pdfplumber(file_bytes, filename)

    def _parse_pdf_pymupdf(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[ParsedPage] = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("dict")["blocks"]
            page_text_parts: list[str] = []
            headings: list[str] = []

            for block in blocks:
                if block.get("type") != 0:  # 0 = text block
                    continue
                for line in block.get("lines", []):
                    line_text = " ".join(
                        span["text"] for span in line.get("spans", [])
                    ).strip()
                    if not line_text:
                        continue

                    # Detect headings by font size / bold flags
                    spans = line.get("spans", [])
                    if spans:
                        avg_size = sum(s["size"] for s in spans) / len(spans)
                        is_bold = any(s["flags"] & 2**4 for s in spans)  # bold flag
                        if avg_size > 11 or is_bold:
                            headings.append(line_text)

                    page_text_parts.append(line_text)

            page_text = "\n".join(page_text_parts)
            pages.append(ParsedPage(
                page_number=page_num + 1,
                text=page_text,
                headings=headings,
            ))

        full_text = "\n\n".join(p.text for p in pages)
        doc.close()

        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            pages=pages,
            full_text=full_text,
            total_pages=len(pages),
            metadata={"parser": "pymupdf"},
        )

    def _parse_pdf_pdfplumber(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        import pdfplumber

        pages: list[ParsedPage] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                # Extract table text as supplemental
                tables = page.extract_tables()
                table_text = ""
                for table in tables:
                    for row in table:
                        if row:
                            table_text += " | ".join(str(c) for c in row if c) + "\n"

                combined = text + ("\n" + table_text if table_text else "")
                pages.append(ParsedPage(
                    page_number=i + 1,
                    text=combined.strip(),
                    headings=[],
                ))

        full_text = "\n\n".join(p.text for p in pages)
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            pages=pages,
            full_text=full_text,
            total_pages=len(pages),
            metadata={"parser": "pdfplumber"},
        )

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _parse_docx(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs_text: list[str] = []
        headings: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                headings.append(text)
            paragraphs_text.append(text)

        # Include table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs_text.append(row_text)

        full_text = "\n".join(paragraphs_text)

        # DOCX treated as single "page" for metadata consistency
        pages = [ParsedPage(page_number=1, text=full_text, headings=headings)]

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            pages=pages,
            full_text=full_text,
            total_pages=1,
            metadata={"parser": "python-docx", "headings": headings},
        )

    # ── TXT ───────────────────────────────────────────────────────────────────

    def _parse_txt(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        lines = text.splitlines()
        headings = [
            line.strip() for line in lines
            if line.strip() and (line.isupper() or re.match(r"^\d+[\.\)]\s+[A-Z]", line))
        ]

        pages = [ParsedPage(page_number=1, text=text, headings=headings)]
        return ParsedDocument(
            filename=filename,
            file_type="txt",
            pages=pages,
            full_text=text,
            total_pages=1,
            metadata={"parser": "plaintext"},
        )
